import pandas as pd 
import streamlit as st 
import os
import numpy as np
import gspread
from openpyxl import load_workbook
from pathlib import Path
import traceback
from docx import Document
from docx.shared import Inches
from io import BytesIO
import time
from google.oauth2.service_account import Credentials
from oauth2client.service_account import ServiceAccountCredentials
from streamlit_gsheets import GSheetsConnection
from datetime import datetime 
import datetime as dt

if 'but' not in st.session_state:
        st.session_state.but = False
if st.session_state.but:
# if st.session_state.sub and st.session_state.but:
          st.success('**Form has been downloaded, check your downloads**')
          time.sleep(2)
          st.markdown("""
         <meta http-equiv="refresh" content="0">
               """, unsafe_allow_html=True)
cluster = ''
pos = ''
vist = ''
consent = ''
district = ''
facility = ''
partners = ''
age = ''
name = ''
htn = ''
dm = ''
tbaction = ''
lam = ''
AS = ''
MH = ''
sex = ''
others = ''
otherissue = ''
adher = ''
cd = ''
vl = ''
tblam = ''
crag = ''
ellig = ''
chid = ''
tested = ''
pos = ''
linked = ''
post = ''
screened = ''
presumed = ''
picked = ''

st.write('**COMMUNITY CLIENT ENCOUNTER FORM**')

file = r'CLUSTERS.csv'
dfa = pd.read_csv(file)

clusters = dfa['CLUSTER'].unique()

cluster = st.radio('**CHOOSE A CLUSTER**', clusters, index= None, horizontal=True)

if not cluster:
    st.stop()
else:
    pass
if cluster:
    dfd = dfa[dfa['CLUSTER'] == cluster].copy()
    districts = dfd['DISTRICT'].unique()
    district = st.radio('**CHOOSE A district**', districts, index= None, horizontal=True)

if not district:
    st.stop()
else:
    pass
col1,col2, col3 = st.columns([2,1,2])
with col1:
    if district:
        fac = dfa[dfa['DISTRICT'] == district].copy()
        facilities = fac['FACILITY'].unique()
        facility = st.selectbox('**SELECT FACILITY**', facilities, index= None)

if not facility:
    st.stop()
else:
    pass

filen = r'ALL.csv'
dfn = pd.read_csv(filen)

dfna = dfn[dfn['FACILITY'] == facility].copy()

with col3:
    art = st.text_input('**INPUT ART No.**')

if not art:
    st.stop()
else:
    artstr = art
    dfart = pd.DataFrame({'A':'A','ART':[art]})
    dfart['ART'] = dfart['ART'].astype(str)
    dfart['ART'] = dfart['ART'].str.replace('[^0-9]', '', regex=True)
    art = dfart.iloc[0,1]
    try:
       art = int(art)
    except:
        st.warning('CHECK THIS ART NO, IT HAS ONLY LETTERS, NO NUMBERS')
        st.stop()
    dfna['ART-NUM'] = dfna['ART-NUM'].astype(int)
    dfna = dfna[dfna['ART-NUM']==art]
    shpe = dfna.shape[0]
    if shpe ==0:
        st.write(f"**ART No. {artstr} not found, in put this client's VL and date of sample collection**")
        col1, col2 = st.columns(2)
        results = col1.number_input('**RESULTS**', value=None, min_value=0)
        if not results:
            st.stop()
        elif results <200:
            st.info('NS results can not be below 200')
            st.stop()
        else:
            dob = col2.date_input('DATE OF SAMPLE COLLECTION', value=None)
            if not dob:
                st.stop()
            else:
                tod = dt.date.today()
                if dob > tod:
                    st.warning("DATE OF SAMPLE COLLECTION SHOULDN'T BE GREATER THAN TODAY")
                    st.stop()
                else:
                    dus = tod - dob
                    dus = int(dus.days)
                    if dus > 120:
                        due = 'DUE'
                    else:
                        due = 'NOT DUE'
                        st.markdown(f'**{art} WAS BLED ON {dob} AND THE RESULTS WERE {results} AND IS {due} FOR REBLEEDING**')
                pass
    else:  
        dfna['ART'] = dfna['ART'].astype(int)        
        client = dfna[dfna['ART'] ==art].copy()

        dob = client.iloc[0,5]
        results = client.iloc[0,6]
        due = client.iloc[0,7]

        st.markdown(f'**{art} WAS BLED ON {dob} AND THE RESULTS WERE {results} AND IS {due} FOR REBLEEDING**')

col1, col2,col3 = st.columns([2,1,2])
age = col1.number_input('AGE', min_value=0, max_value=None, placeholder='Age in years', value=None)
if not age:
    st.stop()
sex = col3.radio('GENDER', options=['Female', 'Male'],index=None, horizontal=True)
if sex =='Female':
    if age>14:
        col1, col2 = st.columns(2)
        pmtct = ['PREGNANT', 'BREAST FEEDING', 'N/A']
        pm = col1.radio("**WHAT IS THIS CLIENT'S PMTCT STATUS**", options=pmtct, horizontal=True, index=None)
        if not pm:
            st.stop()
    else:
        pm = 'N/A'
else:
    pm = 'N/A'

if sex == 'Female':
    him = 'her'
    she = 'she'
else:
    him = 'his'
    she = 'he'

if not sex:
    st.stop()
elif sex and not age:
    st.warning('**IN PUT THE AGE BEFORE YOU PROCEED**')
else:
    pass
file3 = r'DISTRICT.csv'
add = pd.read_csv(file3)

hod = add['DISTRICT'].unique()
st.write("**CLIENT'S ADDRESS**")

col1, col2 = st.columns(2)
dist = col1.selectbox("CLIENT'S HOME DISTRICT", hod, index=None)
if not dist:
    st.stop()
else:
    pass

vil = col2.text_input(f'**input {him} home village, (where this visit is happening)**')
if not vil:
    st.stop()
else:
    pass
st.write(f'**{art} is a {age} year old {sex} from {dist} district, {vil} Village**')

consent = st.pills('**Have you obtained consent for this session?**', options=['YES', 'NO'])#, index=None, horizontal=True)

if consent =='YES':
    IACs = ['FIRST', 'SECOND', 'THIRD', 'FORTH', 'FIFTH', 'SIXTH']
    IAC = st.pills('**WHICH LEVEL OF IAC IS THIS**', options = IACs)
    if not IAC:
        st.stop()
    else:
        pass
else:
    st.warning("YOU CAN'T PROCEED WITHOUT CONSENT")
    st.stop()
st.divider()
st.write('**PILL COUNT, DETERMINE THE NUMBER OF PILLS:**')
col1, col2 = st.columns(2)
with col1:
    had = st.number_input('**SWALLOWED IN ONE MONTH**', min_value=0,  value=None)
with col2:
    supposed = st.number_input('**SUPPOSED TO BE SWALLOWED IN ONE MONTH**', min_value=0, value=None)
if not had:
    st.stop()
if not supposed:
    st.stop()
if not had and supposed:
    st.stop()
elif had and supposed:
    if had > supposed:
        st.warning('THIS CLIENT TOOK MORE PILLS THAN THEY WERE SUPPOSED TO TAKE ?, CHECK AGAIN !!')
        st.stop()
    else:
        adher = round(int(had)/int(supposed) *100)

        if adher >94:
            st.write(f'**THIS IS AN ADHERENCE SCORE OF {adher} % WHICH IS GOOD**')
            ad = 'GOOD'
        elif adher > 84:
            st.write(f'**THIS IS AN ADHERENCE SCORE OF {adher} % WHICH IS FAIR**')
            ad = 'FAIR'
        else:
            st.write(f'**THIS IS AN ADHERENCE SCORE OF {adher} % WHICH IS POOR**')
            ad = 'POOR'
st.divider()

st.write('**BARRIERS TO SUPPRESSION (CLICK ON ALL THAT ARE APPLICABLE)**')
 
social = ['NON-DISCLOSURE','STIGMA AND DISCRIMINATION', 'DYSFUNCTIONAL FAMILY SUPPORT', 'SUBSTANCE ABUSE (eg ALCOHOL, TOBACCO)',
            'LACK OF LIFE SURVIVAL SKILLS', 'MULTIPLE SEXUAL PARTNERS','NONE' ]

econ = ['POVERTY', 'NO PERMANENT HOME ADDRESS', 'HAS ONE OR NO MEAL A DAY','PARENTS ARE POOR (if child)', 'NO ECONOMIC ISSUE']

health = ['MALNOURISHED', 'OTHER CORMORBIDITIES (eg TB, NCDs)', 'NO OTHER DISEASES']

psych = ['DENIAL/COPING WITH RESULTS', 'ANGER/STRESS MANAGEMENT', 'BEREAVEMENT', 'FEAR/ANXIETY', 'NO ISSUE']

spir = ['SPIRITUAL ISSUES', 'MISSES DRUGS/APPOINTMENTS', 'NO SPIRITUAL ISSUE']

socials = []
col1, col2 = st.columns(2)
with col1:
    st.write('**SOCIAL ISSUES**')
    for option in social:
        if st.checkbox(option):#,  key =f'{option}')
             socials.append(option)


with col2:
    st.write('**ECONOMICAL ISSUES**')
    econs = []
    for option in econ:
       if st.checkbox(option):#, key =f'{option}')
            econs.append(option)
  
if len(socials) ==0:
    st.warning('**IF NO SOCIAL ISSUE, SELECT NONE**')
    st.stop()
else:
    pass

if len(socials) >1:
    if 'NONE' in socials:
        st.warning("**YOU CAN'T SEECT NONE WHEN YOU HAVE CHOSEN ANOTHER SOCIAL ISSUE**")
        st.info('**UNCHECK NO SOCIAL ISSUE OR SELECT IT ALONE TO PROCEED**')
        st.stop()
else:
    pass
if len(econs) ==0:
    st.warning('**IF NO ECONOMIC ISSUE, SELECT NO ECONOMIC ISSUE**')
    st.stop()
else:
    pass
if len(econs) >1:
    if 'NO ECONOMIC ISSUE' in econs:
        st.warning("**YOU CAN'T SELECT NO ECONOMIC ISSUE WHEN YOU HAVE CHOSEN ANOTHER ECONOMIC ISSUE**")
        st.info('**UNCHECK NO ECONOMIC ISSUE OR SELECT IT ALONE TO PROCEED**')
        st.stop()
else:
    pass
st.divider()
col1, col2 = st.columns(2)
with col1:
    st.write('**HEALTH ISSUES**')

    healths = []
    for option in health:
        if st.checkbox(option):#,  key =f'{option}')
            healths.append(option)
with col2:
    st.write('**PSYCHOLOGICAL ISSUES**')

    psychs = []
    for option in psych:
        if st.checkbox(option):#,  key =f'{option}')
            psychs.append(option)

if len(healths) ==0:
    st.warning('**IF NO HEALTH ISSUE, SELECT NO OTHER DISEASES**')
    st.stop()
else:
    pass

if len(healths) >1:
    if 'NO OTHER DISEASES' in healths:
        st.warning("**YOU CAN'T SELECT NO OTHER DISEASES WHEN YOU HAVE CHOSEN OTHER HEALTHL ISSUE**")
        st.info('**UNCHECK NO HEALTH ISSUE OR SELECT IT ALONE TO PROCEED**')
        st.stop()
else:
    pass
if len(healths)>0:
    if 'OTHER CORMORBIDITIES (eg TB, NCDs)' in healths:
        otherdiz = col1.text_input('**SPECIFY THE OTHER DISEASE**', placeholder='eg Diabetes, HTN, TB, Stroke')

if len(psychs) ==0:
    st.warning('**IF NO PSYCHOLOGICAL ISSUE, SELECT NO ISSUE**')
    st.stop()
else:
    pass

if len(psychs) >1:
    if 'NO ISSUE' in psychs:
        st.warning("**YOU CAN'T SELECT NO ISSUE WHEN YOU HAVE CHOSEN ANOTHER PSYCHOLOGICAL ISSUE**")
        st.info('**UNCHECK NO ISSUE OR SELECT IT ALONE TO PROCEED**')
        st.stop()
else:
    pass


st.divider()

col1, col2 = st.columns(2)
with col1:
    st.write('**SPIRITUAL ISSUES**')

    spirs = []
    for option in spir:
       if st.checkbox(option):#,  key =f'{option}')
            spirs.append(option)

if len(spirs) ==0:
    st.warning('**IF NO SPIRITUAL ISSUE, SELECT NO SPIRITUAL ISSUE**')
    st.stop()
else:
    pass

if len(spirs) >1:
    if 'NO SPIRITUAL ISSUE' in spirs:
        st.warning("**YOU CAN'T SELECT NO SPIRITUAL ISSUE WHEN YOU HAVE CHOSEN ANOTHER SPIRITUAL ISSUE**")
        st.info('**UNCHECK NO SPIRITUAL ISSUE OR SELECT IT ALONE TO PROCEED**')
        st.stop()
else:
    pass

st.divider()
others = st.radio('**ANY OTHER BARRIER IDENTIFIED BUT NOT CAPTURED ABOVE**', options =['YES', 'NO'], index= None, horizontal=True)
otherissue = ''
if not others:
    st.stop()
elif others =='NO':
    pass
elif others =='YES':
    otherissue = st.text_area('**Write here any other issue not highlighted above**')

if others =='YES':
    if not otherissue:
        st.stop()

st.divider()
act = st.text_area('**WHAT ACTIONS HAVE YOU AGREED UP ON TO ADDRESS THE BARRIERS ABOVE?**')
if not act:
    st.stop()
st.divider()
st.write('**NCD SCREENING (CHOOSE ONE OPTION FOR EACH, OR CHOOSE NOT DONE)**')
if age <15:
    st.info('NCD SCREENING STARTS AT 15 YRS, IF THIS CHILD IS DIABETIC, INCLUDE IT IN COMORBITIES UNDER HEALTH ISSUES')
if age >14:
    htncodes = ['1', '2', '3','4','5', '6' ,7,8,'NOT DONE']
    DMcodes = ['1', '2', '3','4','5', '6',7,8 ,'NOT DONE']
    AScodes = ['1', '2', '3','4','5', '6', 7,8,'NOT DONE']
    MHcodes = ['1', '2', '3','4','5', '6', 7,8,'NOT DONE']

    col1, col2, =st.columns(2)

    htn = col1.selectbox('**HTN SCREENING**', htncodes, index=None)
    dm = col2.selectbox('**DIABETES SCREENING**', DMcodes, index=None)
    AS = col1.selectbox('**ALCOHOL/SUBSTANCE ABUSE SCREENING**', AScodes, index=None)
    MH = col2.selectbox('**MENTAL HEALTH SCREENING**', MHcodes, index=None)
    one = r'1.jpg'
    if not htn:
        col,col2 = st.columns(2)
        col1.write('**No option chosen for HTN screening**')
        with col2.expander('**CLICK HERE TO SEE THE CODES**'):
            st.image(r'1.jpg', caption='-')

        st.stop()
    elif not dm:
        col,col2 = st.columns(2)
        col1.write('**No option chosen for DM screening**')
        with col2.expander('**CLICK HERE TO SEE THE CODES**'):
            st.image(r'2.jpg', caption='-')
        st.stop()
    elif not AS:
        col,col2 = st.columns(2)
        col1.write('**No option chosen for Alcohol/Substance abuse**')
        with col2.expander('**CLICK HERE TO SEE THE CODES**'):
            st.image(r'1.jpg', caption='-')
        st.stop()
    elif not MH:
        col,col2 = st.columns(2)
        col1.write('**No option chosen for Mental Healthe assesment**')
        with col2.expander('**CLICK HERE TO SEE THE CODES**'):
            st.image(r'2.jpg', caption='-')
        st.stop()
    if AS == 'NOT DONE':
        if 'SUBSTANCE ABUSE (eg ALCOHOL, TOBACCO)' in socials:
            st.write('**Among social issues, you selected substance abuse, so this client needs ALOCOHOL/SUBSTANCE ABUSE screening**')

    if MH == 'NOT DONE':
        if len(psychs)>1:
            if 'NO ISSUE' not in psychs:
                st.write('**Yo selected some psychological issues above, so this client needs a Mental Health screening**')
    if htn in ['2', '3','4','5', '6']:
        st.warning('THIS CLIENT SHOULD BE REFERRED FOR HTN MANAGEMENT')
    if dm in ['2', '3','4','5', '6']:
        st.warning('THIS CLIENT SHOULD BE REFERRED FOR DIABETES MANAGEMENT')
    if MH in ['2', '3','4','5', '6']:
        st.warning('THIS CLIENT SHOULD BE REFERRED FOR MENTAL HEALTH MANAGEMENT')
    if AS in ['2', '3','4','5', '6']:
        st.warning('THIS CLIENT SHOULD BE REFERRED FOR  SUBSTANCE ABUSE MANAGEMENT')

else:
    pass

#st.write('**CD-4 TESTING (CLICK ON WHAT IS APPLICABLE)**')
st.divider()
if age <5:
    cd4s = ['CD4 SAMPLE PICKED', 'WAS BLED ALREADY', 'LLV CLIENT']
    st.info('VISITECT IS NOT USED IN CHILDREN <5 YRS, IF NOT ALREADY BLED, PICK A SAMPLE FOR CD4 TESTING')
else:
    cd4s = ['CD4 SAMPLE PICKED', 'WAS BLED ALREADY', 'VISITECT USED', 'LLV CLIENT']

cd = st. pills('**CD-4 TESTING (CLICK ON WHAT IS APPLICABLE)**',options= cd4s)
if not cd:
    st.stop()

if cd =='CD4 SAMPLE PICKED':
    st.info('The picked sampe should be followed up and an update be made in the update section of this form')
elif cd =='VISITECT USED':
    vist =st.radio('**VISTECT RESULTS**', options =['ABOVE REFERENCE', 'BELOW REFERENCE'], horizontal=True, index=None)
    if age < 10:
         if  vist == 'BELOW REFERENCE':
            st.info('TB LAM AND SERUM CRAG ARE NOT DONE IN CHILDREN <10 YEARS, FOR AHD SCREENING, USE SYMPTOMATIC SCREENING LIKE ICF GUIDE, OR PICK A SMAPLE FOR CRP')
         else:
             pass
    else:
        if  vist == 'BELOW REFERENCE':
            st.write('CLIENT NEEDS AHD SCREENING (TB, CrAg)')
            col1,col2 = st.columns(2)
            lam = col1.radio('**DO YOU HAVE A URINE TB-LAM KIT**', options = ['YES', 'NO'], horizontal=True, index=None)
            if lam == 'YES':
                    tblam = col2.radio('**TBLAM RESULTS**', options=['POS','NEG'], horizontal=True, index=None)
                    if not tblam:
                        st.stop()
                    if tblam =='POS':
                        col1.warning('MUST START ANTI TBs')
                        tbaction = col2.radio('**TREATMENT GIVEN**', options=['STARTED ON ANTI-TBS', 'TO BE FOLLOWED UP'], index=None, horizontal=True)
                        if not tbaction:
                            st.stop()
            elif lam=='NO':
                st.success('PICK A URINE SAMPLE FOR TB LAM TESTING,WILL BE UPDATED LATER')
            else:
                st.stop()
        if  vist == 'BELOW REFERENCE':
            st.write('PICK A SAMPE FOR SERUM CRAG')
            col1,col2 = st.columns(2)
            crag = col1.radio('**SAMPLE FOR SERUM CRAG PICKED?**', options = ['YES', 'NO'], horizontal=True, index=None)
            if crag == 'YES':
                    col2.info('**THIS WILL NEED AN UPDATE LATER**')
            elif crag=='NO':
                col2.success('PICK A SAMPLE OR REFER FOR SERUM CrAg, WILL BE FOLLOWED UP')
            else:
                st.stop()
else:
    pass
if not cd:
    st.stop()
st.divider()
col1,col2, col3 = st.columns([1,1,1])
col2.write('**APN SECTION**')
if age <15:
    st.write(f'**APN STARTS AT 15 YEARS, THIS ONE IS A {age} YEAR OLD WHO IS NOT ELLIGIBLE**')
    st.write('GO TO THE NEXT SECTION')
elif age>14:
    if 'MULTIPLE SEXUAL PARTNERS' in socials:
        st.write('YOU SELECTED MULTIPLE SEXUAL PARTNERS, IN SOCIAL ISSUES, SO ELICIT SOME PARTNERS')
    col1,col2 = st.columns(2)
    partners = col1.number_input('**No. OF SEXUAL PARTNERS ELICITED**', min_value=0, value=None)
    if partners or partners == 0:
        if partners >0:
            st.markdown(f'**An update on the partners elicited will be needed, note them in the APN register**')
        elif partners ==0:
            pass
    elif not partners:
        st.stop()
        
st.divider()
st.write('**INDEX BIOLOGICAL TESTING**')
ellig = st.radio('**ARE THERE ELLIGIBLE CHILDREN FOR TESTING IN THE HOUSEHOLD**', options = ['YES', 'NO'], horizontal=True, index=None)
if not ellig:
    st.stop()
elif ellig == 'NO':
    pass
elif ellig == 'YES':
    col1, col2 = st.columns(2)
    chid = col1.number_input('**HOW MANY?**',min_value=0, value=None)
    if chid ==0:
        st.warning("**Elligible children can't be zero, choose NO instead**")
        st.stop()
    if not chid:
        st.stop()
    col1,col2 = st.columns(2)
    tested = col1.number_input(f'**OF THE {chid}, HOW MANY HAVE YOU TESTED**', min_value=0, value=None)
    if tested or  tested ==0:
        if tested > chid:
            st.warning("YOU CAN'T TEST MORE THAN THOSE ELLIGIBLE")
            st.stop()
        elif tested ==0:
            pass
        if tested == 1:
            pos = col2.radio('**WAS HE/SHE POS?**', options =['YES', 'NO'], horizontal=True, index=None)
            if not pos:
                st.stop()
            elif pos =='YES':
                pos = 1
            elif pos == 'NO':
                linked = 0 
                pos = 0
        elif tested >1:
            post = col2.number_input(f'**OF THE {tested}, HOW MANY ARE POS**', min_value=0, value=None)
            if post or post == 0:   
                if post == 0:
                    pass
                elif post>0:
                    if post > tested:
                        st.warning("**THE POS CAN'T BE MORE THAN THOSE TESTED**")
                        st.stop()
                    col1, col2 = st.columns(2)    
                    linked = col1.number_input(f'**TOTAL LINKED TO CARE**', min_value=0, value=None)
                    if linked or linked==0:
                        if linked>=0 :
                            if linked > post:
                                st.warning("**YOU CAN'T LINK MORE THAN THOSE TESTED**")
                                st.stop()
                            notlinked = post - linked
                            if notlinked >0:
                                if notlinked ==1:
                                    st.warning(f'{notlinked} client has not been linked, they have to be followed up')
                                elif notlinked >1:
                                    st.warning(f'{notlinked} clients have not been linked, they have to be followed up')
                            else:
                                pass 
                    else:
                        st.stop()
            else:
                st.stop()
    else:
        st.stop()
st.divider()    
col1,col2,col3 = st.columns([1,2,1])

col2.write('**TB SCREENING**')
col1,col2 = st.columns(2)

screened = col1.number_input('**TOTAL No. SCREENED FOR TB**', min_value=0, value=None)
if screened or screened ==0:
    if screened ==0:
        pass
    else:
        presumed =col2.number_input('**TOTAL PRESUMED FOR TB**', min_value=0, value=None)
        if presumed or presumed ==0:
            if presumed ==0:
                pass
            elif presumed > screened:
                st.warning("**PRESUMED CAN'T BE GREATER THAN SCREENED**")
                st.stop()
            else:
                col1, col2 = st.columns(2)
                picked =col1.number_input('**TOTAL SPUTUM SAMPLES PICKED**', min_value=0, value=None)
                if picked or picked ==0:
                    if picked ==0:
                        st.warning('**ALL PRESUMED CASES SHOULD ACCESS A TB TEST, THIS HAS TO BE FOLLOWED UP**')
                        pass
                    elif picked > presumed:
                        st.warning("**YOU CAN'T PICK MORE SAMPLES THAN THOSE PRESUMED**")
                        st.stop()
                    else:
                        st.success('PICKED SAMPLES WILL BE FOLLOWED UP AND AN UPDATE MADE IN THE UPDATE SECTION')
                else:
                    st.stop()
        else:
            st.stop()
else:
    st.stop()
st.divider()

st.write('**OTHER PREVENTION SERVICES GIVEN TO THIS CLIENT OR MEMBERS OF THE HOUSEHOLD**')

col1,col2 = st.columns(2)
prev = ['CONDOMS','REFERED FOR VMMC', 'PREP SERVICES', 'GBV SCREENING','NO SERVICE']
condoms = ''
vmmc = ''
prepsc = ''
prepel = ''
prepnum = ''

prevs = []
for option in prev:
    if col1.checkbox(option, key=option):
        prevs.append(option)
if 'CONDOMS' in prevs:
    condoms = col2.number_input('NUMBER OF CONDOM PIECES GIVEN', min_value=0, value=None)
    if not condoms:
        st.stop()
if 'REFERED FOR VMMC' in prevs:
    vmmc = col2.number_input('NUMBER REFERRED FOR VMMC',min_value=0, value=None)
    if not vmmc:
        st.stop()
    else:
        col2.info('These will be followed up for circumcision')

if 'PREP SERVICES' in prevs:
    st.write("**complete the PREP cascade, or unclick it if you didn't offer it**")
    col1,col2,col3 = st.columns(3)
    prepsc = col1.number_input('**Number screened for PREP**', value=None, step=1)
    if not prepsc:
        st.stop()
    else:
        prepel = col2.number_input('**Number elligible for PREP**', value=None, step=1)
        if prepel or prepel ==0:
            if prepel == 0:
                pass
            elif prepel > prepsc:
                st.warning("**Number elligible can't be greater than number screened**")
                st.stop()
            elif prepel:
                prepnum = col3.number_input('**Number initiated on PREP**', value=None, step=1)
                if prepnum or prepnum ==0:
                    if prepnum > prepel:
                        st.warning("**Number initiated can't be greater than number elligible**")
                        st.stop()
                    else:
                        pass
                else:
                    st.stop()
            else:
                st.stop()
        else:
            st.stop()
    
if len(prevs)==0:
    st.info('CLICK ON NO SERVICE IF NONE WAS GIVEN')
    st.stop()
if len(prevs)>1:
    if 'NO SERVICE' in prevs:
        st.info('YOU CAN NOT CHOOSE NO SERVICE TOGETHER WITH OTHER OPTIONS')
        st.stop()
st.divider()
col1,col2,col3 = st.columns([1,2,1])
col2.write('**ECONOMIC SERVICES**')
col1,col2 = st.columns(2)
st.write('**ECONOMIC SERVICES ADVISED TO THE CLIENT OR HOUSE HOLD MEMBERS**')
economies = ['VSLA (Village Savings and Loans Association)', 'PDM', 'ADVISED ON BACKYARD GARDENING', 'NO ADVICE']

econis = []
for option in economies:
    if st.checkbox(option, key=option):
        econis.append(option)
if len(econis)==0:
    st.info('CLICK ON NO ADVICE IF NONE WAS GIVEN')
    st.stop()
if len(econis)>1:
    if 'NO ADVICE' in econis:
        st.info('YOU CAN NOT CHOOSE NO ADVICE TOGETHER WITH OTHER OPTIONS')
        st.stop()

st.divider()
col1,col2,col3 = st.columns([1,2,1])
col2.write('**VL SECTION**')
reason = ''
if due == 'DUE':
    st.info('**WHEN REBLEEDING NS, TWO PLASMA SAMPLES SHOULD BE PICKED, MAKE SURE YOU REQUEST FOR A DR TEST**')
    vl = st.radio('**HAS THE CLIENT BEEN REBLED FOR VL?**',options = ['YES', 'NO'], horizontal=True, index=None)
    if not vl:
        st.stop()
    elif vl == 'YES':
        st.write('Results will be followed up')
    elif vl =='NO':
        reason = st.text_input('**STATE ANY REASON FOR NOT REBLEEDING THE CLIENT**')
else:
    col1,col2 = st.columns([1,8])
    col2.write('**CLIENT IS NOT YET DUE FOR REBLEEDING**')

st.divider()
st.write('**TO CAPTURE CORDINATES, YOU NEED TO SWITCH ON YOUR LOCATION**')
allow = st.radio('**IS YOUR LOCATION ON**', options =['YES', 'NO'], index=None, horizontal = True)

if not allow:
    st.stop()
elif allow =='NO':
    st.warning("** YOU CAN'T PROCEED WITHOUT LOCATION**")
    st.stop()
elif allow == 'YES':
        st.write('**COPY THESE CORDINATES AND PASTE THEM BELOW**')
            # HTML/JavaScript component for getting browser geolocation
        geolocation_html = """
            <script>
            function getLocation() {
                if (navigator.geolocation) {
                    navigator.geolocation.getCurrentPosition(showPosition, showError);
                } else {
                    alert("Geolocation is not supported by this browser.");
                }
            }
            
            function showPosition(position) {
                const lat = position.coords.latitude;
                const lon = position.coords.longitude;
                document.getElementById("output").value = `${lat},${lon}`;
            }
            
            function showError(error) {
                alert("Switch on your location.");
            }
            
            getLocation();
            </script>
            <input id="output" type="text" readonly>
            <button onclick="getLocation()">Refresh Location</button>
            """

        # geolocation_html = """
        # <script>
        # function getLocation() {
        #     if (navigator.geolocation) {
        #         navigator.geolocation.getCurrentPosition(showPosition, showError);
        #     } else {
        #         alert("Geolocation is not supported by this browser.");
        #     }
        # }
        
        # function showPosition(position) {
        #     const lat = position.coords.latitude;
        #     const lon = position.coords.longitude;
        #     document.getElementById("output").value = `${lat},${lon}`;
        # }
        
        # function showError(error) {
        #     alert("Switch on your location.");
        # }
        
        # getLocation();
        # </script>
        # <input id="output" type="text" readonly>
        # """
        
        st.write('**HERE IS YOUR CURRENT LOCATION, COPY THE CORDINATES AND PASTE THEM BELOW**')
        st.warning('**May take some time to generate them, please wait**')
        st.components.v1.html(geolocation_html, height=100)
        
        # # Inform users about the accuracy
        # st.write("Note: Using browser geolocation provides more precise results compared to IP-based services.")
col1,col2 = st.columns([1,2])
cords = col1.text_input('**PASTE THE ABOVE CORDINATES HERE**')
if not cords:
    st.stop()
else:
    pass                
st.divider()
col1,col2 = st.columns([1,2])
name = col1.text_input('**Name of the HW who did this visit**')
if not name:
    st.stop()
else:
    pass 
name2 = col2.text_input('**Name of the CHW for this NS**')
if not name2:
    st.stop()
else:
    pass 
st.divider()
col1,col2,col3 = st.columns([1,1,2])

col3.write('**SUMMARY**')

st.markdown(f'**You have offered {IAC} IAC to {artstr}, a {age} year old {sex} from {dist}, {vil} village. {she} was bled on {dob}, results were {results} copies/mL and is {due} for rebleeding**')
st.markdown(' ')
st.markdown(f'**PMTCT STATUS : {pm}**')
st.markdown(f'**ADHERENCE SCORE : {adher} % which is {ad}**')
st.divider()
col1,col2 = st.columns([1,8])
col2.write('**BARRIERS TO SUPPRESSION IDENTIFIED ARE:**')
socialx = ','.join(socials)
psychx = ','.join(psychs)
econx = ','.join(econs)
spirx = ','.join(spirs)
healthx = ','.join(healths)
col1,col2 = st.columns(2)
col1.markdown(f'**SOCIAL ISSUES : {socialx}**')
col2.markdown(f'**PSYCHOLOGICAL ISSUES : {psychx}**')
col1.markdown(f'**ECONOMICAL ISSUES : {econx}**')
col2.markdown(f'**HEALTH ISSUES : {healthx}**')
col1.markdown(f'**SPIRITUAL ISSUES : {spirx}**')
st.divider()
if others =='YES':
    st.markdown(f'**OTHER ISSUES IDENTIFIED ARE: {otherissue}**')
    #st.divider()
else:
    pass
st.markdown(f'**YOU HAVE AGREED ON: {act}**')
st.divider()
if age > 14:
    col1,col2, col3 = st.columns([1,2,1])
    col2.write('**NCD SCREENING CODES**')
    col1,col2 = st.columns([1,1])
    col1.markdown(f'**HTN : {htn}**')
    col2.markdown(f'**DM : {dm}**')
    col1.markdown(f'**MH : {MH}**')
    col2.markdown(f'**AS : {AS}**')
    st.divider()
    # col1,col2 = st.columns([2,1])
    # col1.write(f'**You ellicited {partners} partners who are to be followed up**')
if ellig == 'YES':
    if chid ==1:
        st.write(f'In addition, there was {chid} elligible child, you tested {tested}, and found {post} positive and hence linked {linked} to care')
    else:
        st.write(f'In addition, there were {chid} elligible children, you tested {tested}, and found {post} positive and hence linked {linked} to care')
else:
    st.write('There were no elligible children for index biological')
if age > 14:
    if partners >0:
        st.write(f'**You also elicited {partners} partners, who need to be followed up**')
    else:
        st.write(f'**You also elicited no partners**')
else:
    pass
if screened >0:
    st.write(f'**Yo screened {screened} for TB, prseumed {presumed} and picked {picked} sample(s) that should be followed up**')
else:
    st.write(f'**Yo did not screen any for TB**')

if len(prevs) ==1:
    if 'NO SERVICE' in prevs:
        prevs = []
    else:
        pass
if len   (econis)==1:
    if 'NO ADVICE' in econis:
        econis = []
    else:
        pass
otherprev = econis + prevs
prevx = ','.join(prevs)
econix = ','.join(econis)

if len(otherprev) ==0:
    pass
else:
    otherprev = ','.join(otherprev)
    st.write(f'**Other services provided are: {otherprev}**')

col1, col2, col3 = st.columns(3)
submit = col3.button('**SUBMIT**')
dob = str(dob)
todi = str(tod)
row1 = [cluster, district, facility,art, results, dob, age, sex, pm, dist, vil, cords, IAC, adher,ad, htn,dm,AS, MH]
row2 = [facility,art,socialx, econx, healthx, psychx, spirx, otherissue, act, prevx, condoms, vmmc, econix,vl, reason, name, name2, todi]
row3 = [facility,art,cd, vist, lam, tblam, tbaction, crag, partners, ellig, chid, tested, pos, linked, post, screened, presumed, picked]
if 'sub' not in st.session_state:
    st.session_state.sub = False
if not submit:
    st.session_state.sub = False
    st.stop()
else:
    st.success(f'THANK YOU {name}')
    secrets = st.secrets["connections"]["gsheets"]
        # Prepare the credentials dictionary
    credentials_info = {
            "type": secrets["type"],
            "project_id": secrets["project_id"],
            "private_key_id": secrets["private_key_id"],
            "private_key": secrets["private_key"],
            "client_email": secrets["client_email"],
            "client_id": secrets["client_id"],
            "auth_uri": secrets["auth_uri"],
            "token_uri": secrets["token_uri"],
            "auth_provider_x509_cert_url": secrets["auth_provider_x509_cert_url"],
            "client_x509_cert_url": secrets["client_x509_cert_url"]
        }
            
    try:
        # Define the scopes needed for your application
        scopes = ["https://www.googleapis.com/auth/spreadsheets",
                "https://www.googleapis.com/auth/drive"]
         
        credentials = Credentials.from_service_account_info(credentials_info, scopes=scopes)
            
            # Authorize and access Google Sheets
        client = gspread.authorize(credentials)
            
            # Open the Google Sheet by URL
        spreadsheetu = "https://docs.google.com/spreadsheets/d/1qGCvtnYZ9SOva5YqztSX7wjh8JLF0QRw-zbX9djQBWo"
        spreadsheet = client.open_by_url(spreadsheetu)
        sheet1 = spreadsheet.worksheet("DEMO")
        sheet2 = spreadsheet.worksheet("ISSUES")
        sheet3 = spreadsheet.worksheet("TESTS")
        sheet1.append_row(row1, value_input_option='RAW')
        sheet2.append_row(row2, value_input_option='RAW')
        sheet3.append_row(row3, value_input_option='RAW')
        st.session_state.sub = True
        # st.markdown("""
        #      <meta http-equiv="refresh" content="0">
        #        """, unsafe_allow_html=True)
        
        if st.session_state.sub:
          st.info('**Download this form before form refreshes**')
          def create_docx():
                   
               sp = ''
               barx = socialx + econx + healthx + psychx + spirx + otherissue
               bar = barx
              
               document = Document()
               document.add_heading ('  IDI MWP CLIENT ENCOUNTER FORM', 0)
               P = document.add_paragraph(f'   {sp} {sp} This visit, conducted on {todi} was a {IAC} IAC session for client {art}, a {int(age)} year old {sex} from {vil} village, {dist} district, located at {cords}')
               table1 = document.add_table(rows=1,cols=3, style='Table Grid')
               table1.cell(0,0).text = f'RESULTS: {results} copies/mL'
               table1.cell(0,1).text = f'BLED ON: {dob}'
               table1.cell(0,2).text = f'PMTCT: {pm}'
               p = document.add_paragraph('')
               p = document.add_paragraph('')
               p.add_run('ADHRENCE BARRIERS, ACTIONS AGREED UP ON AND SERVICES GIVEN').bold=True
               table1 = document.add_table(rows=5,cols=2, style='Table Grid')
               table1.cell(0,0).text = 'ADHERENCE SCORE:'
               table1.cell(0,1).text = f'{adher} % ({ad})'
               table1.cell(1,0).text = 'BARRIERS IDENTIFIED:'
               table1.cell(1,1).text = bar
               table1.cell(2,0).text = 'ACTIONS AGREED UPON:'
               table1.cell(2,1).text = str(act)
               table1.cell(3,0).text = 'PREVENTION SERVICES GIVEN:'
               table1.cell(3,1).text = str(prevx)
               table1.cell(4,0).text = 'ECONOMIC ADVICE:'
               table1.cell(4,1).text = str(econix)
               for row in table1.rows:
                   row.cells[0].width = Inches(1.5)
                   row.cells[1].width = Inches(5)
               p = document.add_paragraph('')
               p = document.add_paragraph('')
               p.add_run('TESTS DONE').bold=True
               table2 = document.add_table(rows=3,cols=2, style='Table Grid')
               table2.cell(0,0).text = 'REBLED FOR VL?:'
               table2.cell(0,1).text = f'{vl}'
               table2.cell(1,0).text = 'CD4 TESTING:'
               table2.cell(1,1).text = f'{cd}'
               table2.cell(2,0).text = 'SCREENED FOR TB:'
               table2.cell(2,1).text = f'{(screened)}'
               p = document.add_paragraph('')
               p = document.add_paragraph('')
               p.add_run('APN, NCD CODES').bold=True
               table2 = document.add_table(rows=1,cols=5, style='Table Grid')
               table2.cell(0,0).text = f'Partners: {partners}'
               table2.cell(0,1).text = f'HTN: {htn}'
               table2.cell(0,2).text = f'DM: {dm}'
               table2.cell(0,3).text = f'MH: {MH}'
               table2.cell(0,4).text = f'AS: {AS}'
               p = document.add_paragraph('')
               p = document.add_paragraph('')
               p.add_run('Name of H/worker:').bold =True
               p.add_run(f'{name}').bold =True
               p = document.add_paragraph('')
               p.add_run('Name of CHW:').bold =True
               p.add_run(f'{name2}').bold =True
      
               doc_io = BytesIO()
               document.save(doc_io)
               doc_io.seek(0)  # Move pointer to the start of the file
               return doc_io
          
          doc_file = create_docx()
             # Provide a download button
          but = st.download_button(
                 label=f"DOWNLOAD FORM FOR ART NO: {art:,.0f}",
                 data=doc_file,
                 file_name=f"FORM FOR ART NO: {art:,.0f}.docx",
                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
          if but:
              st.session_state.but = True

          # if st.session_state.sub and st.session_state.but:
          #     st.success('**Form has been downloaded, check your downloads**')
          #     time.sleep(2)
          #     st.markdown("""
          #    <meta http-equiv="refresh" content="0">
          #          """, unsafe_allow_html=True)
          elif not but:
              st.info('**DOWNLOAD THIS FORM BEFORE THIS PAGE REFRESHES**')
              time.sleep(40)
              st.markdown("""
                 <meta http-equiv="refresh" content="0">
                   """, unsafe_allow_html=True)
    except Exception as e:
            # Log the error message
        st.session_state.sub = False
        st.write(f"CHECK: {e}")
        st.write(traceback.format_exc())
        st.write("** POOR NETWORK, COULDN'T CONNECT TO GOOGLE SHEET, SUBMIT AGAIN**")
        st.stop()




