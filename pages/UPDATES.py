import pandas as pd 
import streamlit as st 
import os
import numpy as np
import gspread
from openpyxl import load_workbook
from pathlib import Path
import traceback
import time
from google.oauth2.service_account import Credentials
from oauth2client.service_account import ServiceAccountCredentials
from streamlit_gsheets import GSheetsConnection
from datetime import datetime 
import datetime as dt
from docx import Document
from docx.shared import Inches
from io import BytesIO

cluster = ''
consent = ''
district = ''
facility = ''
partners = 0
age = ''
name = ''
htn = ''
dm = ''
AS = ''
MH = ''
sex = ''
others = ''
otherissue = ''
adher = ''
alread =0
notif = ''
pos = ''
neg =''
alread = ''
linked = ''
recent = ''
cd4 = ''
cd4results = ''
tblamdone = ''
tblamres = ''
tblamrx = ''
crag = ''
crares = ''
ccmres = ''
csf = ''
tbsamples = ''
tbtest = ''
tbrest = ''
tbtreat = ''
tbneg = ''
sup = ''


# st.write('**FOLLOW UP SECTION ON AREAS NOT COMPLETED FROM THE FIELD**')
# if 'tx' not in st.session_state:     
#      try:
#         #cola,colb= st.columns(2)
#         conn = st.connection('gsheets', type=GSheetsConnection)
#         exist = conn.read(worksheet= 'DEMO', usecols=list(range(19)),ttl=5)
#         tx = exist.dropna(how='all')
#         st.session_state.tx = tx
#      except:
#          st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
#          st.stop()
# dfdemo = st.session_state.tx.copy()

# if 'txa' not in st.session_state:     
#      try:
#         #cola,colb= st.columns(2)
#         conn = st.connection('gsheets', type=GSheetsConnection)
#         exist = conn.read(worksheet= 'ISSUES', usecols=list(range(18)),ttl=5)
#         txa = exist.dropna(how='all')
#         st.session_state.txa = txa
#      except:
#          st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
#          st.stop()
# dfiss = st.session_state.txa.copy()
# #########################################################################################################

# if 'txb' not in st.session_state:     
#      try:
#         #cola,colb= st.columns(2)
#         conn = st.connection('gsheets', type=GSheetsConnection)
#         exist = conn.read(worksheet= 'TESTS', usecols=list(range(18)),ttl=5)
#         txb = exist.dropna(how='all')
#         st.session_state.txb = txb
#      except:
#          st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
#          st.stop()
# dftest = st.session_state.txb.copy()
###############################################################################################################
if 'tx' not in st.session_state:     
     try:
        #cola,colb= st.columns(2)
        conn = st.connection('gsheets', type=GSheetsConnection)
        exist = conn.read(worksheet= 'DEMO', usecols=list(range(20)),ttl=5)
        tx = exist.dropna(how='all')
        st.session_state.tx = tx
     except:
         st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
         st.stop()
dfdemo = st.session_state.tx.copy()

dfdemo['FACILITY'] = dfdemo['FACILITY'].astype(str)
dfdemo['DT'] = dfdemo['DATE'].astype(str)
facilities = dfdemo['FACILITY'].unique()

dfdemoz = []
for facil in facilities:
     dfissa = dfdemo[dfdemo['FACILITY']==facil].copy()
     dfissa[['YEAR','MONTH','DAY']] = dfissa['DT'].str.split('-', expand=True)
     dfissa[['ART NO', 'YEAR', 'MONTH', 'DAY']] = dfissa[['ART NO', 'YEAR', 'MONTH', 'DAY']].apply(pd.to_numeric, errors='coerce')
     dfissa = dfissa.sort_values(by = ['YEAR', 'MONTH', 'DAY'], ascending = [False, False, False])
     dfissa = dfissa.drop_duplicates(subset = ['ART NO', 'YEAR', 'MONTH', 'DAY'])
     dfdemoz.append(dfissa)
dfdemo = pd.concat(dfdemoz)
dfdemo[['ART NO', 'DAY', 'MONTH']] = dfdemo[['ART NO', 'DAY', 'MONTH']].astype(str)
dfdemo['ART'] = dfdemo['MONTH'] + dfdemo['DAY'] + dfdemo['ART NO'] 


if 'txa' not in st.session_state:     
     try:
        #cola,colb= st.columns(2)
        conn = st.connection('gsheets', type=GSheetsConnection)
        exist = conn.read(worksheet= 'ISSUES', usecols=list(range(18)),ttl=5)
        txa = exist.dropna(how='all')
        st.session_state.txa = txa
     except:
         st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
         st.stop()
dfiss = st.session_state.txa.copy()

dfiss['FACILITY'] = dfiss['FACILITY'].astype(str)
dfiss['DT'] = dfiss['DATE'].astype(str)
facilities = dfiss['FACILITY'].unique()

dfissz = []
for facil in facilities:
     dfissa = dfiss[dfiss['FACILITY']==facil].copy()
     dfdemu = dfdemo[dfdemo['FACILITY'] == facil].copy()
     try:
          distr = dfdemu.iloc[0,1]
          clus = dfdemu.iloc[0,0]
     except:
           distr = ''
           clus = ''
     dfissa[['YEAR','MONTH','DAY']] = dfissa['DT'].str.split('-', expand=True)
     dfissa[['ART NO', 'YEAR', 'MONTH', 'DAY']] = dfissa[['ART NO', 'YEAR', 'MONTH', 'DAY']].apply(pd.to_numeric, errors='coerce')
     dfissa = dfissa.sort_values(by = ['YEAR', 'MONTH', 'DAY'], ascending = [False, False, False])
     dfissa = dfissa.drop_duplicates(subset = ['ART NO', 'YEAR', 'MONTH', 'DAY'])
     dfissa['DISTRICT'] = distr
     dfissa['CLUSTER'] = clus
     dfissz.append(dfissa)
dfiss = pd.concat(dfissz)
dfiss[['ART NO', 'DAY', 'MONTH']] = dfiss[['ART NO', 'DAY', 'MONTH']].astype(str)
dfiss['ART'] = dfiss['MONTH'] + dfiss['DAY'] + dfiss['ART NO'] 


#########################################################################################################

if 'txb' not in st.session_state:     
     try:
        #cola,colb= st.columns(2)
        conn = st.connection('gsheets', type=GSheetsConnection)
        exist = conn.read(worksheet= 'TESTS', usecols=list(range(19)),ttl=5)
        txb = exist.dropna(how='all')
        st.session_state.txb = txb
     except:
         st.write("POOR NETWORK, COULDN'T CONNECT TO DELIVERY DATABASE")
         st.stop()
dftest = st.session_state.txb.copy()

dftest['FACILITY'] = dftest['FACILITY'].astype(str)
dftest['DT'] = dftest['DATE'].astype(str)
facilities = dftest['FACILITY'].unique()

dftestz = []
for facil in facilities:
     dftisa = dftest[dftest['FACILITY']==facil].copy()
     dfdemu = dfdemo[dfdemo['FACILITY'] == facil].copy()
     try:
          distr = dfdemu.iloc[0,1]
          clus = dfdemu.iloc[0,0]
     except:
           distr = ''
           clus = ''
     dftisa[['YEAR','MONTH','DAY']] = dftisa['DT'].str.split('-', expand=True)
     dftisa[['ART NO', 'YEAR', 'MONTH', 'DAY']] = dftisa[['ART NO', 'YEAR', 'MONTH', 'DAY']].apply(pd.to_numeric, errors='coerce')
     dftisa = dftisa.sort_values(by = ['YEAR', 'MONTH', 'DAY'], ascending = [False, False, False])
     dftisa = dftisa.drop_duplicates(subset = ['ART NO', 'YEAR', 'MONTH', 'DAY'])
     dftisa['DISTRICT'] = distr
     dftisa['CLUSTER'] = clus
     dftestz.append(dftisa)
     
dftest = pd.concat(dftestz)
dftest[['DAY', 'MONTH']] = dftest[['DAY', 'MONTH']].astype(int)
dftest[['ART NO', 'DAY', 'MONTH']] = dftest[['ART NO', 'DAY', 'MONTH']].astype(str)
dftest['ART'] = dftest['MONTH'] + dftest['DAY'] + dftest['ART NO'] 
################################################################################################################
#ONE MERGED ONE
facilities = dfdemo['FACILITY'].unique()

dfdemz = []
dfissx = dfiss.drop(columns = ['DT', 'DISTRICT', 'CLUSTER', 'YEAR', 'MONTH', 'DAY', 'ART NO', 'DATE'])
for fac in facilities:
     dfdemu = dfdemo[dfdemo['FACILITY']==fac].copy()
     dfissu = dfissx[dfissx['FACILITY']==fac].copy()
     dfissu = dfissu.drop(columns ='FACILITY')
     dfdemu['ART'] = pd.to_numeric(dfdemu['ART'], errors = 'coerce')
     dfissu['ART'] = pd.to_numeric(dfissu['ART'], errors = 'coerce')
     dfd = pd.merge(dfdemu, dfissu, on = 'ART', how = 'inner')
     dfdemz.append(dfd)
     
dfdemo1 = pd.concat(dfdemz)

############################################################################
#last merged

dfdemoz = []

dfissy = dftest.drop(columns = ['DT', 'DISTRICT', 'CLUSTER', 'YEAR', 'MONTH', 'DAY', 'ART NO', 'DATE'])

for fac in facilities:
     dfdemux = dfdemo1[dfdemo1['FACILITY']==fac].copy()
     dfissh = dfissy[dfissy['FACILITY']==fac].copy()
     dfissh = dfissh.drop(columns ='FACILITY')
     dfdemux['ART'] = pd.to_numeric(dfdemux['ART'], errors = 'coerce')
     dfissh['ART'] = pd.to_numeric(dfissh['ART'], errors = 'coerce') 
     dfd = pd.merge(dfdemux, dfissh, on = 'ART', how = 'inner')
     dfdemoz.append(dfd)

dfdemo2 = pd.concat(dfdemoz)
################################################################################################################

file = r'CLUSTERS.csv'
dfa = pd.read_csv(file)

clusters = dfa['CLUSTER'].unique()

cluster = st.radio('**CHOOSE A CLUSTER**', clusters, index= None, horizontal=True)

if not cluster:
    st.stop()
else:
    pass
if cluster:
    dfd = dfa[dfa['CLUSTER'] == cluster].copy()
    districts = dfd['DISTRICT'].unique()
    district = st.radio('**CHOOSE A district**', districts, index= None, horizontal=True)

if not district:
    st.stop()
else:
    pass
col1,col2, col3 = st.columns([2,1,2])
with col1:
    if district:
        fac = dfa[dfa['DISTRICT'] == district].copy()
        facilities = fac['FACILITY'].unique()
        facility = st.selectbox('**SELECT FACILITY**', facilities, index= None)

if not facility:
    st.stop()
else:
    pass
    
dftest['FACILITY'] = dftest['FACILITY'].astype(str)
dftest = dftest[dftest['FACILITY'] == facility].copy()

dfiss['FACILITY'] = dfiss['FACILITY'].astype(str)
dfiss = dfiss[dfiss['FACILITY'] == facility].copy()

dfdemo['FACILITY'] = dfdemo['FACILITY'].astype(str)
dfdemo = dfdemo[dfdemo['FACILITY'] == facility].copy()
factz = dfdemo['FACILITY'].unique()
num = dfdemo.shape[0]
if num ==0:
     st.warning('**THERE IS NO DATA FOR THIS FACILITY**')
     st.stop()
dfdemoz =[]
for facx in factz:
     dfdemof = dfdemo[dfdemo['FACILITY']==facx].copy()
     dfdemof = dfdemof.drop_duplicates(subset = 'ART NO', keep='last')
     dfdemoz.append(dfdemof)
dfdemo = pd.concat(dfdemoz)

dfissuez =[]
for facx in factz:
     dfissf = dfiss[dfiss['FACILITY']==facx].copy()
     dfissf = dfissf.drop_duplicates(subset = 'ART NO', keep='last')
     dfissuez.append(dfissf)
dfiss = pd.concat(dfissuez)

dftestz =[]
for facx in factz:
     dftestf = dftest[dftest['FACILITY']==facx].copy()
     dftestf = dftestf.drop_duplicates(subset = 'ART NO', keep='last')
     dftestz.append(dftestf)
dftest = pd.concat(dftestz)


filen = r'ALL.csv'
dfn = pd.read_csv(filen)

dfna = dfn[dfn['FACILITY'] == facility].copy()

check = st.pills('**WHAT DO YOU WANT TO DO?**', options = ['CHECK UPDATE STATUS', '','','MAKE UPDATES','', 'DOWNLOAD FORM'])
if 'form' not in st.session_state:
     st.session_state.form = False
if not check:
    st.stop()
elif check == 'MAKE UPDATES':
    col1, col2,col3 = st.columns(3)
    ssss
    art = col1.number_input('**SEARCH ART No.**', value=None, step=1, key = 1)
    if not art:
         st.stop()
    dfdemo2 = dfdemo2[dfdemo2['ART NO'] == art].copy()
    if dfdemo2.shape[0] == 0:
         st.info(f'**ART NO {art} NOT FOUND IN THE DATA BASE**')
         st.stop()
    elif dfdemo2.shape[0] > 1:
         st.info('**THIS CLIENT WAS REVISITED, UPDATES FOR WHICH IAC SESSION ARE YOU MAKING?**')
         options = dfdemoz['IAC'].unique()
         iac = st.select_box('**CHOOSE FROM HERE**', options, index = None)
         if not iac:
              st.stop()
         else:
              dfdemo2 = dfdemo2[dfdemo2['IAC'] == iac].copy()
              dfdemo = dfdemo2[['CLUSTER', 'DISTRICT', 'FACILITY', 'ART NO', 'RESULTS', 'DOB','AGE', 'SEX', 'PMTCT', 'DISTRI', 'VILLAGE', 'CORDS', 'IAC', 'ADH', 'AD','htn', 'dm', 'AS', 'MH']].copy()
              dftest = dfdemo2[['FACILITY', 'ART NO', 'SOCIALS', 'ECONS', 'HEALTH', 'PSYCH','SPIRS', 'OTHERISSUES', 'ACT', 'PREVS', 'CONDOMS', 'VMMC', 'ECONIS','VL', 'REASON', 'NAME', 'NAME2', 'DATE']].copy()
              dfiss  = dfdemo2[['FACILITY', 'ART NO', 'CD4', 'VISITECT', 'LAM', 'TBLAM','TB RX', 'CRAG', 'PARTNERS', 'ELLIG', 'CHILD', 'TESTED', 'POS','LINKED', 'POST', 'SCREENED', 'PRESUMED', 'PICKED']].copy()
    else:
         dfdemo = dfdemo2[['CLUSTER', 'DISTRICT', 'FACILITY', 'ART NO', 'RESULTS', 'DOB','AGE', 'SEX', 'PMTCT', 'DISTRI', 'VILLAGE', 'CORDS', 'IAC', 'ADH', 'AD','htn', 'dm', 'AS', 'MH']].copy()
         dftest = dfdemo2[['FACILITY', 'ART NO', 'SOCIALS', 'ECONS', 'HEALTH', 'PSYCH','SPIRS', 'OTHERISSUES', 'ACT', 'PREVS', 'CONDOMS', 'VMMC', 'ECONIS','VL', 'REASON', 'NAME', 'NAME2', 'DATE']].copy()
         dfiss  = dfdemo2[['FACILITY', 'ART NO', 'CD4', 'VISITECT', 'LAM', 'TBLAM','TB RX', 'CRAG', 'PARTNERS', 'ELLIG', 'CHILD', 'TESTED', 'POS','LINKED', 'POST', 'SCREENED', 'PRESUMED', 'PICKED']].copy()
    dftest= dftest[dftest['ART NO'] == art].copy()
    #dftest['PARTNERS'] = pd.to_numeric(dftest['PARTNERS'], errors = 'coerce')
    dfiss = dfiss[dfiss['ART NO'] == art].copy()
    st.write('**APN SECTION**')
    partners = dftest.iloc[0,8]
    if partners > 0:
         if partners ==1:
              st.write(f'**Client with ART NO {art} had {partners:,.0f} partner ellicited**')
              col1,col2,col3 = st.columns(3)
              notif = col1.radio('**HAS HE/SHE BEEN NOTIFIED**', options=['YES', 'NOT YET', 'UPDATE ALREADY MADE'], horizontal= True, index=None)
              if not notif:
                   st.stop()
              elif notif == 'NOT YET':
                   st.warning('**KINDLY FOLLOW UP ON THIS PARTNER PLEASE**')
              elif notif == 'UPDATE ALREADY MADE':
                   pass
              elif notif == 'YES':
                   notif = 1
                   testapn = col2.radio('**WAS SHE/HE TESTED**', options=['YES', 'NOT YET', 'UPDATE ALREADY MADE'], horizontal= True, index=None)
                   if not testapn:
                        st.stop()
                   elif testapn == 'NOT YET':
                        st.warning('**ENSURE THIS CLIENT IS TESTED**')
                        pass
                   elif testapn == 'UPDATE ALREADY MADE':
                        pass
                   elif testapn == 'YES':
                        tested = 1
                        posapn = col3.radio('**WHAT WAS THE RESULT**', options=['NEG', 'POS', 'KNOWN POS'], horizontal= True, index=None)
                        if not posapn:
                             st.stop()
                        elif posapn in ['NEG', 'KNOWN POS']:
                             if posapn == 'NEG':
                                  neg = 1
                             elif posapn == 'KNOWN POS':
                                  alread = 1
                             pass
                        elif posapn =='POS':
                             pos = 1
                             col1, col2 = st.columns([1,2])
                             linkedapn = col1.radio('**WAS HE/SHE LINKED TO CARE**', options=['YES', 'NO'], horizontal= True, index=None)
                             if linkedapn == 'YES':
                                  linked = 1
                             if not linkedapn:
                                  st.stop()
                             else:
                                  pass
                             recent = col2.radio('**WHAT WERE THE RECENCY RESULTS**', options=['RECENT', 'LONG TERM', 'NOT DONE'], index=None, horizontal=True)
                             if not recent:
                                  st.stop()
                             else:
                                  pass
         if partners > 1:
              st.write(f'**Client with ART NO {art} had {partners:,.0f} partners ellicited**')
              updateparts =  col2.radio('**HAS NOTIFICATION BEEN DONE**', options=['YES', 'NOT YET', 'UPDATE ALREADY MADE'], horizontal= True, index=None)
              if not updateparts:
                   st.stop()
              elif updateparts == 'NOT YET':
                   pass
              elif updateparts == 'UPDATE ALREADY MADE':
                   pass
              elif updateparts == 'YES':
                   st.write(f'**OF THESE {partners:,.0f} partners, how many have been:**')
                   col1,col2, col3 = st.columns([2,1,2])
                   notif = col1.number_input('**NOTIFIED**', value=None, step=1, key=2)
                   if not notif:
                        st.stop()
                   if notif > partners:
                        st.warning("**YOU CAN'T NOTIFY MORE THAN THOSE ELLICITED**")
                   else:
                        col1,col2 = st.columns(2)
                        updatetest =  col2.radio('**HAS TESTING BEEN DONE**', options=['YES', 'NOT YET', 'UPDATE ALREADY MADE', 'NONE ELLIGIBLE'], horizontal= True, index=None)
                        if not updatetest:
                             st.stop()
                        elif updatetest in ['UPDATE ALREADY MADE', 'NONE ELLIGIBLE']:
                             pass
                        elif updatetest == 'NOT YET':
                             st.warning('**PARTNERS SHOULD BE FOLLOWED UP FOR TESTING**')
                        elif updatetest == 'YES':
                             col1,col2,col3 = st.columns(3)
                             tested = col3.number_input('**HOW MANY WERE TESTED**', value=None, step=1, key=3)
                             if tested or tested ==0:
                                       pass
                             else:
                                  st.stop()
                             if notif == tested:
                                  pass
                             elif notif > tested:
                                  col1,col2,col3 = st.columns(3)
                                  alread = col1.number_input('**KNOWN POSITIVE, (input 0 if none)**', value=None, step=1, key=4)
                                  if alread or alread ==0:
                                       pass
                                  else:
                                       st.stop()
                                  if alread or alread ==0:
                                       checkm = tested + alread
                                       if checkm > notif:
                                            st.warning("**TESTED AND KNOWN POS ARE MORE THAN THOSE NOTIFIED**")
                                            st.stop()
                             if not tested:
                                  st.stop()
                             elif tested > notif:
                                  st.warning("**YOU CAN'T TEST MORE THAN THOSE NOTIFIED**")
                             elif tested == 0:
                                  pass
                             else:
                                  st.write('**TEST RESULTS**')
                                  col1,col2, col3 = st.columns(3) 
                                  neg = col1.number_input('**NUMBER NEGATIVE**', value=None, step=1, key = 5)
                                  pos = col2.number_input('**NEWLY POSTIVE**', value=None, step=1, key=6)
                                  if pos or pos ==0:
                                       pass
                                  else:
                                       st.stop()
                                  if neg or neg ==0:
                                       pass
                                  else:
                                       st.warning('Total negative is required or put a 0 (zero)**')
                                       st.stop()
                                  if (neg + pos) > tested:
                                       st.warning('**Number positive and negative is greater than number tested**')
                                       st.stop()
                                  if pos:
                                       linked = col3.number_input('**HOW MANY WERE LINKED**', value=None, step=1, key=7)
                                       if linked or linked ==0:
                                            if linked > pos:
                                                 st.warning("**YOU CAN'T LINK MORE THAN THOSE TESTED**")
                                                 st.stop()
                                       elif not linked:
                                            st.stop()
                                       else:
                                            pass
                                  if pos:
                                       st.write(f'**RECENCY TESTING**')
                                       col1,col2 = st.columns(2)
                                       recent = col1.number_input('**NUMBER WITH RECENT RESULT**', value=None, step=1, key=9)
                                       long = col2.number_input('**NUMBER WITH LONGTERM RESULTS**', value=None, step=1, key=10)
                                       if recent or recent ==0:
                                             pass
                                       else:
                                            st.stop()
                                       if long or long == 0:
                                             pass
                                       else:
                                             st.stop()
                                       if long or long ==0:
                                            if recent or recent ==0:
                                               checki = recent + long
                                               if checki > pos:
                                                   st.warning('**TOTAL WITH RECENCY RESULTS IS GREATER THAN POSITIVES**')
                                                   st.stop()
                         
    ###TB
    cd4 = dftest.iloc[0,2]
    
    if cd4 !='CD4 SAMPLE PICKED':
         st.info('**NO CD4 UPDATES ARE NEEDED, PROCEED TO TB SECTION**')
    elif cd4 == 'CD4 SAMPLE PICKED':
         updcd4 = st.radio('**A CD4 SAMPLE WAS PICKED, DO YOU WANT TO UPDATE THE RESULTS**', options =['YES', 'NOT YET DONE', 'UPDATE WAS ALREADY MADE'],horizontal=True, index=None)
         if not updcd4:
              st.stop()
         elif updcd4 =='NOT YET DONE':
              pass
         elif updcd4 == 'UPDATE WAS ALREADY MADE':
              pass
         else: 
              cd4results = st.radio('**CD4 RESULTS**', options = ['BELOW 200', 'ABOVE 200', 'BELOW REFRCE', 'ABOVE REFRCE'], horizontal=True, index=None)
              if not cd4results:
                   st.stop()
              elif cd4results in ['ABOVE 200', 'ABOVE REFRCE']:
                   pass
              elif cd4results in ['BELOW 200', 'BELOW REFRCE']:
                   tblamdone = st.radio('**WAS TB LAM DONE**', options = ['YES', 'NO'], horizontal=True, index=None)
                   if not tblamdone:
                        st.stop()
                   elif tblamdone == 'NO':
                        st.warning('**CLIENT NEEDS A TB LAM, MAKE SURE YOU UPDATE IT**')
                   else:
                        tblamres = st.radio('**TB LAM RESULTS**', options =['POS', 'NEG'],  horizontal=True, index=None)
                        if not tblamres:
                             st.stop()
                        elif tblamres =='NEG':
                             pass
                        else:
                             tblamrx = st.radio("**WAS THE CLIENT STARTED ON ANTI-TB's**", options = ['YES', 'NO'], horizontal=True, index=None)
                             if not tblamrx:
                                  st.stop()
                             elif tblamrx =='NO':
                                  st.warning('CLIENTSHOULD BE STARTED ON ANTI TBs')
                             else:
                                  pass
                   crag = st.radio('**WAS CRAG DONE**', options = ['YES', 'NO'], horizontal=True, index=None)
                   if not crag:
                        st.stop()
                   elif crag == 'NO':
                        st.warning('**THIS CLIENT NEEDS A SERUM CRAG, MAKE SURE YOU UPDATE IT**')
                   elif crag == 'YES':
                        crares = st.radio('**CRAG RESULTS**', options =['POS', 'NEG'],  horizontal=True, index=None)
                        if not crares:
                             st.stop()
                        elif crares == 'NEG':
                             pass
                        elif crares == 'POS':
                             ccmres = st.radio('**WAS CSF CRAG DONE**', options = ['YES', 'NO', 'CLIENT WAS REFERED'], horizontal=True, index=None)
                             if not ccmres:
                                  st.stop()
                             elif ccmres == 'CLIENT WAS REFERED':
                                  st.warning('**FOLLOW UP ON THIS CLIENT TO ASCERTAIN THEIR CSF CRAG RESULTS**')
                             elif ccmres == 'NO':
                                  st.warning('**CLIENT NEEDS A CSF CRAG, KINDLY FOLLOW UP**')
                             elif ccmres == 'YES':
                                  csf = st.radio('**CSF CRAG RESULTS**', options =['POS', 'NEG'],  horizontal=True, index=None)
                                  if not csf:
                                       st.stop()
                                  elif csf == 'NEG':
                                       st.info('**START THE CLIENT ON HIGH DOSE FLUCONAZOLE**')
                                  elif csf == 'POS':
                                       st.info('**START THE CLIENT ON CCM TREATMENT**')
    tbsamples = dftest.iloc[0,17]
       
    if tbsamples >0:
         st.write('**TB SAMPLES PICKED**')
         if tbsamples==1:
              tbtest  = st.radio('**1 SPUTUM SAMPLE WAS PICKED HAS IT BEEN TESTED**', options = ['YES', 'NOT YET', 'UPDATE ALREADY MADE'], horizontal=True, index=None)
              if not tbtest:
                     st.stop()
              elif  tbtest == 'NOT YET':
                    st.warning('**THE COLLECTED SAMPLE SHOULD BE TESTED**')
                                                                                        
              elif tbtest == 'UPDATE ALREADY MADE':
                    pass
              elif tbtest == 'YES': 
                    tbtest = 1
                    tbrest = st.radio('**WHAT WERE THE RESULTS**', options = ['POS', 'NEG'], horizontal=True, index=None)
                    
                    if not tbrest:
                         st.stop()
                    elif tbrest == 'NEG':
                         tbneg = 1
                         pass
                    elif tbrest == 'POS':
                         tbrest = 1
                         treat = st.radio('**WAS THE CLIENT STARTED ON ANTI TBS**',options = ['YES', 'NO'], horizontal=True, index=None)
                         if not treat:
                              st.stop()
                         elif treat == 'NO':
                              st.warning('**START THIS CLIENT ON ANTI TBs PLEASE**')
                         elif treat == 'YES':
                              tbtreat = 1
                              pass

         if tbsamples>1:
             tbtest  = st.radio(f'**{int(tbsamples)} SPUTUM SAMPLES WERE PICKED, HAVE THEY BEEN TESTED**', options = ['YES', 'NOT YET', 'UPDATE ALREADY MADE'], horizontal=True, index=None)
             if not tbtest:
                  st.stop()
             if tbtest == 'NOT YET':
                  st.warning('**FOLLOW UP TO SEE THAT THEY ARE TESTED**')
             elif tbtest == 'UPDATE ALREADY MADE':
                  pass
             elif tbtest == 'YES':
                  st.write('**HOW MANY WERE (PUT A ZERO WHERE APPLICABLE):**')
                  col1,col2,col3 = st.columns(3)
                  tbrest = col1.number_input('**POSTIVE**', value=None, step=1, key = 'a11')
                  # if tbrest or tbrest==0:
                  #      if tbrest >0:
                  #           tbtreat = col2.number_input('**TREATED**', value=None, step=1, key = 'a13')
                  #           if tbtreat or tbtreat == 0:
                  #                if tbtreat > tbrest:
                  #                     st.warning("TREATED CAN'T BE GREATER THAN THOSE POS")
                  #                     st.stop()
                  #           else:
                  #                st.stop()  
                  # else:
                  #      st.stop()
                  if tbrest or tbrest==0:
                      if tbrest < tbsamples or tbrest==0:
                            col1,col2,col3 = st.columns(3)
                            tbneg = col1.number_input('**NEG**', value=None, step=1, key = 'a12')
                            if tbneg or tbneg == 0:
                                  tbck = tbneg + tbrest
                                  if tbck > tbsamples:
                                      st.warning("THAT'S MORE THAN SAMPLES COLLECTED")
                                      st.stop()
                                  elif tbck < tbsamples:
                                       tbnot = tbsamples -tbck
                                       st.success(f'THIS LEAVES {int(tbnot)} SAMPLE(S) NOT ACCOUNTED FOR')
                            elif not tbneg:
                                 st.stop()
                            else:
                                 pass
                  else: 
                    st.stop()

                  if tbrest or tbrest==0:
                       if tbrest >0:
                            tbtreat = col2.number_input('**TREATED**', value=None, step=1, key = 'a13')
                            if tbtreat or tbtreat == 0:
                                 if tbtreat > tbrest:
                                      st.warning("TREATED CAN'T BE GREATER THAN THOSE POS")
                                      st.stop()
                            else:
                                 st.stop()  
                  else:
                       st.stop()
    else:
         tbsamples = ''
         st.info('**NO SPUTUM SAMPLE WAS PICKED, PROCEED TO VL SECTION**')
    vlsample = dfiss.iloc[0,13]
    
    if vlsample == 'YES': 
         vlrest = st.radio('**A VL SAMPLE WAS PICKED, ARE RESULTS BACK**', options=['YES','NO'], horizontal=True, index=None)
         if not vlrest:
              st.stop()
         elif vlrest =='NO':
              pass
         elif vlrest =='YES':
              col1, col2, col3 = st.columns(3)
              sup = col1.number_input('**WHAT WERE THE RESULTS**',value=None, step=1, key='vl')
              if not sup:
                   st.stop()
              else: 
                   pass
    else:
       # if not vlsample or vlsample == 'NO':
             st.info('**NO VL SAMPLE WAS PICKED, PROCEED TO SUBMIT**')
    col1, col2, col3 = st.columns(3)
    submit = col1.button('SUBMIT')
    
    row1 = [cluster, district, facility, art, partners, notif, pos, neg, alread, linked, recent, cd4, cd4results, tblamdone, tblamres,
                                                    tblamrx, crag, crares, ccmres, csf, tbsamples, tbtest, tbrest, tbtreat, tbneg,sup]
     
    if not submit:
         st.stop()
    else:   
         secrets = st.secrets["connections"]["gsheets"]
             # Prepare the credentials dictionary
         credentials_info = {
                 "type": secrets["type"],
                 "project_id": secrets["project_id"],
                 "private_key_id": secrets["private_key_id"],
                 "private_key": secrets["private_key"],
                 "client_email": secrets["client_email"],
                 "client_id": secrets["client_id"],
                 "auth_uri": secrets["auth_uri"],
                 "token_uri": secrets["token_uri"],
                 "auth_provider_x509_cert_url": secrets["auth_provider_x509_cert_url"],
                 "client_x509_cert_url": secrets["client_x509_cert_url"]
              }
                 
         try:
             # Define the scopes needed for your application
             scopes = ["https://www.googleapis.com/auth/spreadsheets",
                     "https://www.googleapis.com/auth/drive"]
              
             credentials = Credentials.from_service_account_info(credentials_info, scopes=scopes)
                 
                 # Authorize and access Google Sheets
             client = gspread.authorize(credentials)
                 
                 # Open the Google Sheet by URL
             spreadsheetu = "https://docs.google.com/spreadsheets/d/1qGCvtnYZ9SOva5YqztSX7wjh8JLF0QRw-zbX9djQBWo"
             spreadsheet = client.open_by_url(spreadsheetu)
             sheet = spreadsheet.worksheet("UPDATES")
         
             sheet.append_row(row1, value_input_option='RAW')

             st.success(f'THANK YOU, THESE UPDATES HAVE BEEN MADE')
             time.sleep(2)
             st.markdown("""
                  <meta http-equiv="refresh" content="0">
                    """, unsafe_allow_html=True)
         except Exception as e:
                 # Log the error message
             st.info('NOT SUBMITTED')
             st.session_state.sub = False
             st.write(f"CHECK: {e}")
             st.write(traceback.format_exc())
             st.write("** POOR NETWORK, COULDN'T CONNECT TO GOOGLE SHEET, SUBMIT AGAIN**")
             st.stop()
             
         
elif check == 'DOWNLOAD FORM':
     st.session_state.form = True
if st.session_state.form:
    col1, col2,col3 = st.columns(3)
    artu = col1.number_input('**SEARCH ART No.**', value=None, step=1, key =11)
    
    if artu:
         artu = int(float(artu))
         dfdemo['ART NO'] = pd.to_numeric(dfdemo['ART NO'], errors = 'coerce')
         dftest['ART NO'] = pd.to_numeric(dftest['ART NO'], errors = 'coerce')
         dfiss['ART NO'] = pd.to_numeric(dfiss['ART NO'], errors = 'coerce')
         dfdemo = dfdemo[dfdemo['ART NO'] == artu].copy()
         dftest= dftest[dftest['ART NO'] == artu].copy()
         dfiss = dfiss[dfiss['ART NO'] == artu].copy()
    else:
         dfdemo = dfdemo.copy()
         dftest = dftest.copy()
         dfisss = dfiss.copy()
         Arts = dfdemo['ART NO'].unique()

    num = dfdemo.shape[0]
    if num ==0:
         st.info('**No forms available for this facility or ART NO selected**')
         st.stop()
    elif artu or num ==1:
              if not artu:
                   artu = dfdemo.iloc[0,3]
              st.info('**1 form is available for this facility or ART NO selected**')
              def create_docx():
                   dfdemy= dfdemo.copy()
                   dftesty = dftest.copy()
                   dfisy = dfiss.copy()
                   dfdemy['ART NO'] = dfdemy['ART NO'].astype(str)
                   
                   arty = dfdemy.iloc[0,3]
                   
                   a = arty
                   b = dfdemy.iloc[0,12]
                   c = dfdemy.iloc[0,6]
                   d = dfdemy.iloc[0,7]
                   e = dfdemy.iloc[0,10]
                   f = dfdemy.iloc[0,9]
                   g = dfdemy.iloc[0,11]
                   dt = dfisy.iloc[0,17]
                   sp = ''
                   pm = dfdemy.iloc[0,8]
                   dob = dfdemy.iloc[0,5]
                   dr = dfdemy.iloc[0,4]
               
                   a1 = dfisy.iloc[0,2]
                   a2 = dfisy.iloc[0,3]
                   a3 = dfisy.iloc[0,4]
                   a4 = dfisy.iloc[0,5]
                   a5 = dfisy.iloc[0,6]
                   a6 = dfisy.iloc[0,7]
                   bar = f'{a1}, {a2}, {a3}, {a4}, {a5}, {a6}'
                   b1 = dfdemy.iloc[0,13]
                   b2 = dfdemy.iloc[0,14]
                   act = dfisy.iloc[0,8]
                   prev = dfisy.iloc[0,9]
                   econ = dfisy.iloc[0,12]
               
                   cd = dftesty.iloc[0,2]
                   tb = dftesty.iloc[0,14]
                   vl = dfisy.iloc[0,13]
               
                   apn = dftesty.iloc[0,8]
                   ht = dfdemy.iloc[0,15]
                   dm = dfdemy.iloc[0,16]
                   mh = dfdemy.iloc[0,18]
                   As = dfdemy.iloc[0,17]
               
                   hw = dfisy.iloc[0,15]
                   chw = dfisy.iloc[0,16]
                   
                   document = Document()
                   document.add_heading ('  IDI MWP CLIENT ENCOUNTER FORM', 0)
                   P = document.add_paragraph(f'   {sp} {sp} This visit, conducted on {dt} was a {b} IAC session for client {a}, a {c} year old {d} from {e} village, {f} district, located at {g}')
                   table1 = document.add_table(rows=1,cols=3, style='Table Grid')
                   table1.cell(0,0).text = f'RESULTS: {dr} copies/mL'
                   table1.cell(0,1).text = f'BLED ON: {dob}'
                   table1.cell(0,2).text = f'PMTCT: {pm}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('ADHRENCE BARRIERS, ACTIONS AGREED UP ON AND SERVICES GIVEN').bold=True
                   table1 = document.add_table(rows=5,cols=2, style='Table Grid')
                   table1.cell(0,0).text = 'ADHERENCE SCORE:'
                   table1.cell(0,1).text = f'{b1} % ({b2})'
                   table1.cell(1,0).text = 'BARRIERS IDENTIFIED:'
                   table1.cell(1,1).text = bar
                   table1.cell(2,0).text = 'ACTIONS AGREED UPON:'
                   table1.cell(2,1).text = str(act)
                   table1.cell(3,0).text = 'PREVENTION SERVICES GIVEN:'
                   table1.cell(3,1).text = str(prev)
                   table1.cell(4,0).text = 'ECONOMIC ADVICE:'
                   table1.cell(4,1).text = str(econ)
                   for row in table1.rows:
                       row.cells[0].width = Inches(1.5)
                       row.cells[1].width = Inches(5)
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('TESTS DONE').bold=True
                   table2 = document.add_table(rows=3,cols=2, style='Table Grid')
                   table2.cell(0,0).text = 'REBLED FOR VL?:'
                   table2.cell(0,1).text = f'{vl}'
                   table2.cell(1,0).text = 'CD4 TESTING:'
                   table2.cell(1,1).text = f'{cd}'
                   table2.cell(2,0).text = 'TB SCREENING:'
                   table2.cell(2,1).text = f'{tb}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('APN, NCD CODES').bold=True
                   table2 = document.add_table(rows=1,cols=5, style='Table Grid')
                   table2.cell(0,0).text = f'Partners: {apn:,.0f}'
                   table2.cell(0,1).text = f'HTN: {ht}'
                   table2.cell(0,2).text = f'DM: {dm}'
                   table2.cell(0,3).text = f'MH: {mh}'
                   table2.cell(0,4).text = f'AS: {As}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('Name of H/worker:').bold =True
                   p.add_run(f'{hw}').bold =True
                   p = document.add_paragraph('')
                   p.add_run('Name of CHW:').bold =True
                   p.add_run(f'{chw}').bold =True
          
                   doc_io = BytesIO()
                   document.save(doc_io)
                   doc_io.seek(0)  # Move pointer to the start of the file
                   return doc_io
              
              doc_file = create_docx()
     
             # Provide a download button
              st.download_button(
                 label=f"DOWNLOAD FORM FOR ART NO: {artu:,.0f}",
                 data=doc_file,
                 file_name=f"FORM FOR ART NO: {artu:,.0f}.docx",
                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
    else:
         st.info(f'**{num} forms are available for this facility or ART NO selected**')
         for arty in Arts:
              def create_docx():
                   dfdemy = dfdemo[dfdemo['ART NO'] == arty].copy()
                   dfisy = dfiss[dfiss['ART NO'] == arty].copy()
                   dftsty = dftest[dftest['ART NO'] == arty].copy() 
                   a = arty
                   b = dfdemy.iloc[0,12]
                   c = dfdemy.iloc[0,6]
                   d = dfdemy.iloc[0,7]
                   e = dfdemy.iloc[0,10]
                   f = dfdemy.iloc[0,9]
                   g = dfdemy.iloc[0,11]
                   dt = dfisy.iloc[0,17]
                   sp = ''
                   pm = dfdemy.iloc[0,8]
                   dob = dfdemy.iloc[0,5]
                   dr = dfdemy.iloc[0,4]
               
                   a1 = dfisy.iloc[0,2]
                   a2 = dfisy.iloc[0,3]
                   a3 = dfisy.iloc[0,4]
                   a4 = dfisy.iloc[0,5]
                   a5 = dfisy.iloc[0,6]
                   a6 = dfisy.iloc[0,7]
                   bar = f'{a1}, {a2}, {a3}, {a4}, {a5}, {a6}'
                   b1 = dfdemy.iloc[0,13]
                   b2 = dfdemy.iloc[0,14]
                   act = dfisy.iloc[0,8]
                   prev = dfisy.iloc[0,9]
                   econ = dfisy.iloc[0,12]
               
                   cd = dftsty.iloc[0,2]
                   tb = dftsty.iloc[0,14]
                   vl = dfisy.iloc[0,13]
               
                   apn = dftsty.iloc[0,8]
                   ht = dfdemy.iloc[0,15]
                   dm = dfdemy.iloc[0,16]
                   mh = dfdemy.iloc[0,18]
                   As = dfdemy.iloc[0,17]
               
                   hw = dfisy.iloc[0,15]
                   chw = dfisy.iloc[0,16]
                   
                   document = Document()
                   document.add_heading ('  IDI MWP CLIENT ENCOUNTER FORM', 0)
                   P = document.add_paragraph(f'   {sp} {sp} This visit, conducted on {dt} was a {b} IAC session for client {a}, a {c} year old {d} from {e} village, {f} district, located at {g}')
                   table1 = document.add_table(rows=1,cols=3, style='Table Grid')
                   table1.cell(0,0).text = f'RESULTS: {dr} copies/mL'
                   table1.cell(0,1).text = f'BLED ON: {dob}'
                   table1.cell(0,2).text = f'PMTCT: {pm}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('ADHRENCE BARRIERS, ACTIONS AGREED UP ON AND SERVICES GIVEN').bold=True
                   table1 = document.add_table(rows=5,cols=2, style='Table Grid')
                   table1.cell(0,0).text = 'ADHERENCE SCORE:'
                   table1.cell(0,1).text = f'{b1} % ({b2})'
                   table1.cell(1,0).text = 'BARRIERS IDENTIFIED:'
                   table1.cell(1,1).text = bar
                   table1.cell(2,0).text = 'ACTIONS AGREED UPON:'
                   table1.cell(2,1).text = str(act)
                   table1.cell(3,0).text = 'PREVENTION SERVICES GIVEN:'
                   table1.cell(3,1).text = str(prev)
                   table1.cell(4,0).text = 'ECONOMIC ADVICE:'
                   table1.cell(4,1).text = str(econ)
                   for row in table1.rows:
                       row.cells[0].width = Inches(1.5)
                       row.cells[1].width = Inches(5)
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('TESTS DONE').bold=True
                   table2 = document.add_table(rows=3,cols=2, style='Table Grid')
                   table2.cell(0,0).text = 'VL REBLEED:'
                   table2.cell(0,1).text = f'{vl}'
                   table2.cell(1,0).text = 'CD4 TESTING:'
                   table2.cell(1,1).text = f'{cd}'
                   table2.cell(2,0).text = 'TB SCREENING:'
                   table2.cell(2,1).text = f'{tb}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('APN, NCD CODES').bold=True
                   table2 = document.add_table(rows=1,cols=5, style='Table Grid')
                   #apn = int(apn)
                   apn = str(apn)
                   table2.cell(0,0).text = f'Partners: {apn}'
                   table2.cell(0,1).text = f'HTN: {ht}'
                   table2.cell(0,2).text = f'DM: {dm}'
                   table2.cell(0,3).text = f'MH: {mh}'
                   table2.cell(0,4).text = f'AS: {As}'
                   p = document.add_paragraph('')
                   p = document.add_paragraph('')
                   p.add_run('Name of H/worker:').bold =True
                   p.add_run(f'{hw}').bold =True
                   p = document.add_paragraph('')
                   p.add_run('Name of CHW:').bold =True
                   p.add_run(f'{chw}').bold =True
          
                   doc_io = BytesIO()
                   document.save(doc_io)
                   doc_io.seek(0)  # Move pointer to the start of the file
                   return doc_io
              
              doc_file = create_docx()
              
              
             # Provide a download button
              artyz = int(float(arty))
              st.download_button(  
                 label=f"DOWNLOAD FORM FOR ART NO: {artyz:,.0f}",
                 data=doc_file,
                 file_name=f"FORM FOR ART NO: {artyz:,.0f}.docx",
                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                 key = arty
                    )
    
    
elif check == 'MAKE UPDATES':
    pass

st.stop()
